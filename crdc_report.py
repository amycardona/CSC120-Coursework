# generate_crdc_report.py
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os

# -------------------------
# USER CONFIG: change these paths if your files have different names
ENROLL_OVERALL = "Enrollment-Overall (1).xlsx"
ENROLL_EL = "Enrollment-English-Learner (1).xlsx"
ENROLL_504 = "Enrollment-Section-504-only.xlsx"
SUSP_SWD = "One-or-More-OOS-Suspensions-SWD.xlsx"
EXP_SWD = "Expulsions-with-ed-service-SWD.xlsx"

OUT_DIR = "crdc_output"
os.makedirs(OUT_DIR, exist_ok=True)

# -------------------------
# Helper to load the first sheet of a workbook
def load_first_sheet(path):
    xls = pd.ExcelFile(path)
    sheet = xls.sheet_names[0]
    df = pd.read_excel(path, sheet_name=sheet)
    return df

# -------------------------
# Auto-detect numeric column to use as count (heuristic)
def pick_count_column(df):
    # prefer columns with obvious names
    candidates = [c for c in df.columns if any(x in c.lower() for x in ["total", "count", "number", "n_", "students", "susp", "exp"])]
    numeric_cols = df.select_dtypes(include='number').columns.tolist()
    # intersection by name priority
    for c in candidates:
        if c in numeric_cols:
            return c
    # fallback to first numeric column
    if numeric_cols:
        return numeric_cols[0]
    # no numeric column
    return None

# -------------------------
# Load data (catch errors)
def safe_load(path):
    if not os.path.exists(path):
        print(f"File not found: {path}")
        return None
    try:
        return load_first_sheet(path)
    except Exception as e:
        print(f"Error loading {path}: {e}")
        return None

enroll_overall = safe_load(ENROLL_OVERALL)
enroll_el = safe_load(ENROLL_EL)
enroll_504 = safe_load(ENROLL_504)
susp_swd = safe_load(SUSP_SWD)
exp_swd = safe_load(EXP_SWD)

# -------------------------
# Identify columns
enr504_col = pick_count_column(enroll_504) if enroll_504 is not None else None
susp_col = pick_count_column(susp_swd) if susp_swd is not None else None
exp_col = pick_count_column(exp_swd) if exp_swd is not None else None

print("Detected columns (heuristic):")
print("  SWD enrollment column:", enr504_col)
print("  SWD suspensions column:", susp_col)
print("  SWD expulsions column:", exp_col)

# If you prefer an exact column, set it here:
# enr504_col = "Section504_Total"   # <--- edit if needed
# susp_col = "OOS_Suspensions_SWD"  # <--- edit if needed
# exp_col = "Expulsions_SWD"        # <--- edit if needed

# -------------------------
# Compute totals and rates
summary = {}
if enroll_504 is not None and enr504_col:
    summary['swd_enrollment_total'] = int(enroll_504[enr504_col].sum(skipna=True))
else:
    summary['swd_enrollment_total'] = None

if susp_swd is not None and susp_col:
    summary['swd_suspensions_total'] = int(susp_swd[susp_col].sum(skipna=True))
else:
    summary['swd_suspensions_total'] = None

if exp_swd is not None and exp_col:
    summary['swd_expulsions_total'] = int(exp_swd[exp_col].sum(skipna=True))
else:
    summary['swd_expulsions_total'] = None

if summary['swd_enrollment_total'] and summary['swd_suspensions_total'] is not None:
    summary['suspensions_per_100_swd'] = summary['swd_suspensions_total'] / summary['swd_enrollment_total'] * 100
else:
    summary['suspensions_per_100_swd'] = None

if summary['swd_enrollment_total'] and summary['swd_expulsions_total'] is not None:
    summary['expulsions_per_100_swd'] = summary['swd_expulsions_total'] / summary['swd_enrollment_total'] * 100
else:
    summary['expulsions_per_100_swd'] = None

print("Summary computed:", summary)

# -------------------------
# Make simple plots (matplotlib only)
plots = {}

# Plot A: top N rows by suspension count (useful for sample view)
if susp_swd is not None and susp_col:
    top = susp_swd[[susp_col]].fillna(0).sort_values(by=susp_col, ascending=False).head(10)
    fig_path = os.path.join(OUT_DIR, "suspensions_top10.png")
    plt.figure(figsize=(8,4))
    plt.bar(range(len(top)), top[susp_col])
    plt.xticks(range(len(top)), top.index.astype(str), rotation=45, ha='right')
    plt.ylabel(susp_col)
    plt.title("Top 10 rows by suspension count (SWD file)")
    plt.tight_layout()
    plt.savefig(fig_path)
    plt.close()
    plots['suspensions_top10'] = fig_path

# Plot B: top N rows by expulsions
if exp_swd is not None and exp_col:
    top = exp_swd[[exp_col]].fillna(0).sort_values(by=exp_col, ascending=False).head(10)
    fig_path = os.path.join(OUT_DIR, "expulsions_top10.png")
    plt.figure(figsize=(8,4))
    plt.bar(range(len(top)), top[exp_col])
    plt.xticks(range(len(top)), top.index.astype(str), rotation=45, ha='right')
    plt.ylabel(exp_col)
    plt.title("Top 10 rows by expulsions count (SWD file)")
    plt.tight_layout()
    plt.savefig(fig_path)
    plt.close()
    plots['expulsions_top10'] = fig_path

# Plot C: enrollment by race if available in enrollment_overall
if enroll_overall is not None:
    possible_race_cols = [c for c in enroll_overall.columns if any(x in c.lower() for x in ['black', 'white', 'hispanic', 'asian', 'native', 'two or more'])]
    if len(possible_race_cols) >= 2:
        race_sums = enroll_overall[possible_race_cols].sum(skipna=True)
        fig_path = os.path.join(OUT_DIR, "enrollment_race_proportions.png")
        plt.figure(figsize=(6,6))
        plt.pie(race_sums, labels=possible_race_cols, autopct='%1.1f%%')
        plt.title("Enrollment by race (summed across rows)")
        plt.tight_layout()
        plt.savefig(fig_path)
        plt.close()
        plots['enrollment_race'] = fig_path

# -------------------------
# Create Word doc with text and insert plots
doc = Document()
doc.add_heading("Preliminary Report: CRDC 2020-21 — SWD Discipline", level=1)

doc.add_heading("Introduction", level=2)
doc.add_paragraph(
    "This report examines out-of-school suspensions and expulsions for students with disabilities (SWD) using CRDC 2020-21 files."
)

doc.add_heading("Methods", level=2)
doc.add_paragraph(
    "Files used: Enrollment-Overall, Enrollment-English-Learner, Enrollment-Section-504-only, "
    "One-or-More-OOS-Suspensions-SWD, Expulsions-with-ed-service-SWD. "
    "I loaded the first sheet from each Excel file and used a heuristic to pick numeric count columns; you may override the column names at the top of the script."
)

doc.add_heading("Data Summary", level=2)
for k,v in summary.items():
    doc.add_paragraph(f"{k}: {v}")

doc.add_heading("Figures", level=2)
for key, path in plots.items():
    doc.add_paragraph(key)
    doc.add_picture(path, width=Inches(5.5))
    # small caption
    doc.add_paragraph(f"Figure: {key} — source file used and column detected automatically.")

doc.add_heading("References", level=2)
doc.add_paragraph("U.S. Department of Education, Office for Civil Rights. CRDC 2020-21 Estimations (data supplied).")

out_doc = os.path.join(OUT_DIR, "Preliminary_Report_CRDC_SWD.docx")
doc.save(out_doc)
print("Saved Word doc to:", out_doc)
print("Saved plots to:", plots)
